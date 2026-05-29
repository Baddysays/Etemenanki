import json
import platform
import subprocess
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional

import psutil
import requests
from docx import Document
from PyPDF2 import PdfReader
from PySide6.QtCore import QObject, Property, QThread, Signal, Slot, QUrl
from PySide6.QtGui import QGuiApplication
from PySide6.QtQml import QQmlApplicationEngine


APP_NAME = "Etemenanki"
ROOT = Path(__file__).resolve().parents[1]


DEFAULT_MODELS_CATALOG = {
    "models": [
        {"id": "translategemma:4b", "provider": "ollama", "required_ram_gb": 8, "required_vram_gb": 4, "description": "Balanced local translation model."},
        {"id": "translategemma:12b", "provider": "ollama", "required_ram_gb": 16, "required_vram_gb": 10, "description": "Higher quality local model."},
        {"id": "deepseek-chat", "provider": "cloud", "required_ram_gb": 4, "required_vram_gb": None, "description": "Cloud model, no local GPU needed."},
        {"id": "gpt-4.1-mini", "provider": "cloud", "required_ram_gb": 4, "required_vram_gb": None, "description": "Cloud model with strong translation quality."},
    ]
}


def resolve_catalog_path() -> Optional[Path]:
    candidates = [ROOT / "assets" / "models_catalog.json"]
    if getattr(sys, "frozen", False):
        exe_dir = Path(sys.executable).resolve().parent
        candidates += [exe_dir / "assets" / "models_catalog.json", exe_dir / "_internal" / "assets" / "models_catalog.json"]
        meipass = getattr(sys, "_MEIPASS", None)
        if meipass:
            candidates.append(Path(meipass) / "assets" / "models_catalog.json")
    return next((p for p in candidates if p.exists()), None)


def resolve_qml_path() -> Path:
    candidates = [ROOT / "qml" / "Main.qml"]
    if getattr(sys, "frozen", False):
        exe_dir = Path(sys.executable).resolve().parent
        candidates += [exe_dir / "qml" / "Main.qml", exe_dir / "_internal" / "qml" / "Main.qml"]
        meipass = getattr(sys, "_MEIPASS", None)
        if meipass:
            candidates.append(Path(meipass) / "qml" / "Main.qml")
    for c in candidates:
        if c.exists():
            return c
    raise FileNotFoundError("Main.qml not found")


@dataclass
class PcSpecs:
    cpu_name: str
    cpu_cores: int
    ram_gb: float
    vram_gb: Optional[float]


class ModelCatalog:
    def __init__(self, path: Optional[Path]):
        self.data = json.loads(path.read_text(encoding="utf-8")) if path and path.exists() else DEFAULT_MODELS_CATALOG

    def names(self) -> List[str]:
        return [m["id"] for m in self.data["models"]]

    def get(self, model_id: str) -> Dict[str, Any]:
        return next(m for m in self.data["models"] if m["id"] == model_id)


class HardwareInspector:
    @staticmethod
    def get_specs() -> PcSpecs:
        cpu_name = platform.processor() or "Unknown CPU"
        cpu_cores = psutil.cpu_count(logical=False) or psutil.cpu_count() or 1
        ram_gb = round(psutil.virtual_memory().total / (1024**3), 1)
        return PcSpecs(cpu_name, cpu_cores, ram_gb, HardwareInspector._detect_vram())

    @staticmethod
    def _detect_vram() -> Optional[float]:
        try:
            out = subprocess.check_output(
                ["powershell", "-NoProfile", "-Command", "Get-CimInstance Win32_VideoController | Select-Object -ExpandProperty AdapterRAM"],
                stderr=subprocess.DEVNULL,
                text=True,
            )
            vals = [int(v.strip()) for v in out.splitlines() if v.strip().isdigit()]
            return round(max(vals) / (1024**3), 1) if vals else None
        except Exception:
            return None


class AIClient:
    def __init__(self, runtime: str, base_url: str, api_key: str):
        self.runtime = runtime
        self.base_url = base_url.rstrip("/")
        self.api_key = api_key

    def translate(self, model: str, text: str, src: str, dst: str) -> str:
        prompt = f"Translate from {src} to {dst}. Preserve structure. Return only translated text.\n\nText:\n{text}"
        if self.runtime == "local":
            resp = requests.post(
                "http://127.0.0.1:11434/api/chat",
                json={"model": model, "stream": False, "messages": [{"role": "user", "content": prompt}]},
                timeout=240,
            )
            resp.raise_for_status()
            return resp.json().get("message", {}).get("content", "").strip()
        resp = requests.post(
            f"{self.base_url}/chat/completions",
            headers={"Authorization": f"Bearer {self.api_key}", "Content-Type": "application/json"},
            json={"model": model, "messages": [{"role": "user", "content": prompt}], "temperature": 0.2},
            timeout=240,
        )
        resp.raise_for_status()
        return resp.json()["choices"][0]["message"]["content"].strip()


class TranslateWorker(QObject):
    progress = Signal(int, int)
    finished = Signal(str)
    failed = Signal(str)

    def __init__(self, runtime: str, base_url: str, api_key: str, model: str, src: str, dst: str, text: str):
        super().__init__()
        self.runtime = runtime
        self.base_url = base_url
        self.api_key = api_key
        self.model = model
        self.src = src
        self.dst = dst
        self.text = text

    def _chunks(self, text: str, max_chars: int = 2200) -> List[str]:
        chunks: List[str] = []
        i = 0
        while i < len(text):
            j = min(i + max_chars, len(text))
            if j < len(text):
                split = text.rfind("\n", i, j)
                if split > i + 200:
                    j = split
            chunks.append(text[i:j])
            i = j
        return chunks or [""]

    @Slot()
    def run(self) -> None:
        try:
            client = AIClient(self.runtime, self.base_url, self.api_key)
            chunks = self._chunks(self.text)
            out: List[str] = []
            total = len(chunks)
            for idx, ch in enumerate(chunks, start=1):
                out.append(client.translate(self.model, ch, self.src, self.dst))
                self.progress.emit(idx, total)
            self.finished.emit("\n\n".join(out))
        except Exception as exc:
            self.failed.emit(str(exc))


class Backend(QObject):
    modelsChanged = Signal()
    sourceTextChanged = Signal()
    resultTextChanged = Signal()
    statusChanged = Signal()
    progressChanged = Signal()
    pageCountChanged = Signal()
    compatTextChanged = Signal()
    specsBadgeChanged = Signal()

    def __init__(self):
        super().__init__()
        self.catalog = ModelCatalog(resolve_catalog_path())
        self.specs = HardwareInspector.get_specs()
        self._models = self.catalog.names()
        self._source_text = ""
        self._result_text = ""
        self._status = "Ready"
        self._progress = 0
        self._page_count = 1
        self._compat_text = ""
        self._specs_badge = ""
        self.current_file: Optional[Path] = None
        self._thread: Optional[QThread] = None
        self._worker: Optional[TranslateWorker] = None
        self.refresh_compat(self._models[0] if self._models else "")

    @Property("QVariantList", notify=modelsChanged)
    def models(self):
        return self._models

    @Property(str, notify=sourceTextChanged)
    def sourceText(self) -> str:
        return self._source_text

    @Property(str, notify=resultTextChanged)
    def resultText(self) -> str:
        return self._result_text

    @Property(str, notify=statusChanged)
    def status(self) -> str:
        return self._status

    @Property(int, notify=progressChanged)
    def progress(self) -> int:
        return self._progress

    @Property(int, notify=pageCountChanged)
    def pageCount(self) -> int:
        return self._page_count

    @Property(str, notify=compatTextChanged)
    def compatText(self) -> str:
        return self._compat_text

    @Property(str, notify=specsBadgeChanged)
    def specsBadge(self) -> str:
        return self._specs_badge

    def _set_status(self, value: str) -> None:
        self._status = value
        self.statusChanged.emit()

    @Slot(str)
    def refresh_compat(self, model_id: str) -> None:
        if not model_id:
            return
        model = self.catalog.get(model_id)
        req_ram = model.get("required_ram_gb", 0)
        req_vram = model.get("required_vram_gb")
        ok_ram = self.specs.ram_gb >= req_ram
        ok_vram = True if req_vram is None else (self.specs.vram_gb or 0) >= req_vram
        icon = "OK" if (ok_ram and ok_vram) else "WARN"
        self._compat_text = (
            f"Совместимость: {icon}\nCPU: {self.specs.cpu_name} ({self.specs.cpu_cores} cores)\n"
            f"RAM: {self.specs.ram_gb} GB (need {req_ram} GB)\n"
            f"VRAM: {self.specs.vram_gb if self.specs.vram_gb is not None else 'N/A'} GB (need {req_vram if req_vram is not None else 'N/A'} GB)"
        )
        self._specs_badge = (
            f"RAM: {self.specs.ram_gb} / {req_ram} GB\n"
            f"VRAM: {self.specs.vram_gb if self.specs.vram_gb is not None else 'N/A'} / {req_vram if req_vram is not None else 'N/A'} GB"
        )
        self.compatTextChanged.emit()
        self.specsBadgeChanged.emit()

    @Slot(str)
    def chooseFile(self, path_str: str) -> None:
        try:
            path = Path(path_str.replace("file:///", "")).resolve()
            self.current_file = path
            suffix = path.suffix.lower()
            if suffix in {".txt", ".md"}:
                text = path.read_text(encoding="utf-8", errors="ignore")
                pages = 1
            elif suffix == ".docx":
                doc = Document(path)
                text = "\n".join(p.text for p in doc.paragraphs)
                pages = max(1, min(12, len(doc.paragraphs) // 8 + 1))
            elif suffix == ".pdf":
                reader = PdfReader(str(path))
                pages = len(reader.pages)
                text = "\n\n".join((p.extract_text() or "") for p in reader.pages).strip()
            else:
                self._set_status("Unsupported file type")
                return
            self._source_text = text
            self._page_count = max(1, min(pages, 120))
            self.sourceTextChanged.emit()
            self.pageCountChanged.emit()
            self._set_status(f"Loaded: {path.name}")
        except Exception as exc:
            self._set_status(f"Load error: {exc}")

    @Slot(str)
    def setSourceText(self, text: str) -> None:
        self._source_text = text
        self.sourceTextChanged.emit()

    @Slot(str)
    def setResultText(self, text: str) -> None:
        self._result_text = text
        self.resultTextChanged.emit()

    @Slot(str, str, str, str, str, str)
    def startTranslate(self, runtime: str, model: str, src: str, dst: str, base_url: str, api_key: str) -> None:
        if not self._source_text.strip():
            self._set_status("Source is empty")
            return
        if self._thread is not None:
            self._set_status("Translation already running")
            return

        self._progress = 0
        self.progressChanged.emit()
        self._set_status("Starting translation...")

        self._thread = QThread()
        self._worker = TranslateWorker(runtime, base_url, api_key, model, src, dst, self._source_text)
        self._worker.moveToThread(self._thread)
        self._thread.started.connect(self._worker.run)
        self._worker.progress.connect(self._on_progress)
        self._worker.finished.connect(self._on_finished)
        self._worker.failed.connect(self._on_failed)
        self._worker.finished.connect(self._thread.quit)
        self._worker.failed.connect(self._thread.quit)
        self._thread.finished.connect(self._cleanup_thread)
        self._thread.start()

    @Slot(int, int)
    def _on_progress(self, idx: int, total: int) -> None:
        self._progress = int((idx / max(total, 1)) * 100)
        self.progressChanged.emit()
        self._set_status(f"Translating {idx}/{total}")

    @Slot(str)
    def _on_finished(self, text: str) -> None:
        self._result_text = text
        self.resultTextChanged.emit()
        self._progress = 100
        self.progressChanged.emit()
        self._set_status("Translation complete")

    @Slot(str)
    def _on_failed(self, message: str) -> None:
        self._set_status(f"Translate error: {message}")

    @Slot()
    def _cleanup_thread(self) -> None:
        if self._worker is not None:
            self._worker.deleteLater()
        if self._thread is not None:
            self._thread.deleteLater()
        self._worker = None
        self._thread = None

    @Slot(str)
    def saveResult(self, path_str: str) -> None:
        try:
            out = Path(path_str.replace("file:///", "")).resolve()
            text = self._result_text.strip()
            if not text:
                self._set_status("No translation to save")
                return
            if out.suffix.lower() == ".docx":
                doc = Document()
                for line in text.splitlines():
                    doc.add_paragraph(line)
                doc.save(out)
            else:
                out.write_text(text, encoding="utf-8")
            self._set_status(f"Saved: {out}")
        except Exception as exc:
            self._set_status(f"Save error: {exc}")


def main() -> None:
    app = QGuiApplication(sys.argv)
    engine = QQmlApplicationEngine()
    backend = Backend()
    engine.rootContext().setContextProperty("backend", backend)
    engine.load(QUrl.fromLocalFile(str(resolve_qml_path())))
    if not engine.rootObjects():
        sys.exit(1)
    sys.exit(app.exec())


if __name__ == "__main__":
    main()
