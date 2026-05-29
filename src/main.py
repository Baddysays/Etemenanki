import json
import platform
import subprocess
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Callable, Dict, List, Optional, Tuple

import psutil
import requests
from docx import Document
from PyPDF2 import PdfReader
from PySide6.QtCore import QObject, QRunnable, Qt, QThreadPool, Signal
from PySide6.QtWidgets import (
    QApplication,
    QCheckBox,
    QComboBox,
    QFileDialog,
    QFormLayout,
    QFrame,
    QGridLayout,
    QGroupBox,
    QHBoxLayout,
    QLabel,
    QLineEdit,
    QMainWindow,
    QMessageBox,
    QPushButton,
    QPlainTextEdit,
    QProgressBar,
    QListWidget,
    QListWidgetItem,
    QSplitter,
    QVBoxLayout,
    QWidget,
)

APP_NAME = "Etemenanki"
ROOT = Path(__file__).resolve().parents[1]

DEFAULT_MODELS_CATALOG = {
    "models": [
        {"id": "translategemma:4b", "provider": "ollama", "required_ram_gb": 8, "required_vram_gb": 4, "description": "Balanced local translation model."},
        {"id": "translategemma:12b", "provider": "ollama", "required_ram_gb": 16, "required_vram_gb": 10, "description": "Higher quality local model."},
        {"id": "deepseek-chat", "provider": "cloud", "required_ram_gb": 4, "required_vram_gb": None, "description": "Cloud model, no local GPU needed."},
        {"id": "gpt-4.1-mini", "provider": "cloud", "required_ram_gb": 4, "required_vram_gb": None, "description": "Cloud model with strong translation quality."},
    ]
}

LANG = {
    "ru": {"title": "Etemenanki - Переводчик документов", "choose_file": "Выбрать файл", "translate": "Перевести", "save_as": "Сохранить как", "runtime": "Режим ИИ", "model": "Модель", "source": "Исходный текст", "result": "Перевод", "compat": "Совместимость с ПК", "ready": "Готово"},
    "en": {"title": "Etemenanki - Document Translator", "choose_file": "Choose file", "translate": "Translate", "save_as": "Save as", "runtime": "AI runtime", "model": "Model", "source": "Source text", "result": "Translation", "compat": "PC compatibility", "ready": "Ready"},
}


def resolve_catalog_path() -> Optional[Path]:
    candidates = [ROOT / "assets" / "models_catalog.json"]
    if getattr(sys, "frozen", False):
        exe_dir = Path(sys.executable).resolve().parent
        candidates += [exe_dir / "assets" / "models_catalog.json", exe_dir / "_internal" / "assets" / "models_catalog.json"]
        meipass = getattr(sys, "_MEIPASS", None)
        if meipass:
            candidates.append(Path(meipass) / "assets" / "models_catalog.json")
    return next((p for p in candidates if p.exists()), None)


@dataclass
class PcSpecs:
    cpu_name: str
    cpu_cores: int
    ram_gb: float
    vram_gb: Optional[float]


class WorkerSignals(QObject):
    finished = Signal()
    result = Signal(object)
    error = Signal(str)
    progress = Signal(object)


class Worker(QRunnable):
    def __init__(self, fn: Callable[..., Any], *args: Any, **kwargs: Any):
        super().__init__()
        self.fn = fn
        self.args = args
        self.kwargs = kwargs
        self.signals = WorkerSignals()

    def run(self) -> None:
        try:
            self.kwargs["progress_callback"] = self.signals.progress
            result = self.fn(*self.args, **self.kwargs)
            self.signals.result.emit(result)
        except Exception as exc:
            self.signals.error.emit(str(exc))
        finally:
            self.signals.finished.emit()


class HardwareInspector:
    @staticmethod
    def get_specs() -> PcSpecs:
        cpu_name = platform.processor() or "Unknown CPU"
        cpu_cores = psutil.cpu_count(logical=False) or psutil.cpu_count() or 1
        ram_gb = round(psutil.virtual_memory().total / (1024**3), 1)
        return PcSpecs(cpu_name, cpu_cores, ram_gb, HardwareInspector._detect_vram())

    @staticmethod
    def _detect_vram() -> Optional[float]:
        try:
            out = subprocess.check_output(["powershell", "-NoProfile", "-Command", "Get-CimInstance Win32_VideoController | Select-Object -ExpandProperty AdapterRAM"], stderr=subprocess.DEVNULL, text=True)
            vals = [int(v.strip()) for v in out.splitlines() if v.strip().isdigit()]
            return round(max(vals) / (1024**3), 1) if vals else None
        except Exception:
            return None


class ModelCatalog:
    def __init__(self, path: Optional[Path]):
        self.data = json.loads(path.read_text(encoding="utf-8")) if path and path.exists() else DEFAULT_MODELS_CATALOG

    def names(self) -> List[str]:
        return [m["id"] for m in self.data["models"]]

    def get(self, model_id: str) -> Dict[str, Any]:
        return next(m for m in self.data["models"] if m["id"] == model_id)


class AIClient:
    def __init__(self, runtime: str, base_url: str, api_key: str):
        self.runtime = runtime
        self.base_url = base_url.rstrip("/")
        self.api_key = api_key

    def translate(self, model: str, text: str, src: str, dst: str) -> str:
        prompt = f"Translate from {src} to {dst}. Preserve structure. Return only translated text.\n\nText:\n{text}"
        if self.runtime == "local":
            resp = requests.post("http://127.0.0.1:11434/api/chat", json={"model": model, "stream": False, "messages": [{"role": "user", "content": prompt}]}, timeout=240)
            resp.raise_for_status()
            return resp.json().get("message", {}).get("content", "").strip()
        resp = requests.post(
            f"{self.base_url}/chat/completions",
            headers={"Authorization": f"Bearer {self.api_key}", "Content-Type": "application/json"},
            json={"model": model, "messages": [{"role": "user", "content": prompt}], "temperature": 0.2},
            timeout=240,
        )
        resp.raise_for_status()
        return resp.json()["choices"][0]["message"]["content"].strip()


class TranslatorEngine:
    def __init__(self, client: AIClient):
        self.client = client

    @staticmethod
    def chunk_text(text: str, max_chars: int = 2200) -> List[str]:
        chunks, i = [], 0
        while i < len(text):
            j = min(i + max_chars, len(text))
            if j < len(text):
                split = text.rfind("\n", i, j)
                if split > i + 200:
                    j = split
            chunks.append(text[i:j])
            i = j
        return chunks or [""]

    def translate_text(self, model: str, text: str, src: str, dst: str, progress_cb: Optional[Callable[[int, int], None]] = None) -> str:
        chunks = self.chunk_text(text)
        out: List[str] = []
        total = len(chunks)
        for i, chunk in enumerate(chunks, start=1):
            out.append(self.client.translate(model, chunk, src, dst))
            if progress_cb:
                progress_cb(i, total)
        return "\n\n".join(out)

    def extract_pdf(self, path: Path) -> str:
        return "\n\n".join((page.extract_text() or "") for page in PdfReader(str(path)).pages).strip()


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.lang = "ru"
        self.catalog = ModelCatalog(resolve_catalog_path())
        self.specs = HardwareInspector.get_specs()
        self.current_file: Optional[Path] = None
        self.thread_pool = QThreadPool.globalInstance()
        self.result_text = ""
        self.page_count = 0
        self._build_ui()
        self._apply_styles()
        self._refresh_model_info()
        self.set_status(self.tr("ready"))

    def tr(self, key: str) -> str:
        return LANG[self.lang].get(key, key)

    def _build_ui(self) -> None:
        self.setWindowTitle(self.tr("title"))
        self.resize(1460, 900)
        root = QWidget()
        self.setCentralWidget(root)
        page = QVBoxLayout(root)
        page.setContentsMargins(14, 14, 14, 14)
        page.setSpacing(8)

        header = QFrame()
        header.setObjectName("HeaderCard")
        header_row = QHBoxLayout(header)
        header_row.setContentsMargins(12, 8, 12, 8)
        brand = QVBoxLayout()
        brand.setSpacing(0)
        title = QLabel("Etemenanki")
        title.setObjectName("Title")
        subtitle = QLabel("AI-переводчик документов")
        subtitle.setObjectName("Subtitle")
        brand.addWidget(title)
        brand.addWidget(subtitle)
        header_row.addLayout(brand)
        header_row.addStretch(1)
        self.hw_badge = QLabel()
        self.hw_badge.setObjectName("Badge")
        header_row.addWidget(self.hw_badge)
        page.addWidget(header)

        controls = QGroupBox("Панель управления")
        grid = QGridLayout(controls)
        grid.setHorizontalSpacing(10)
        grid.setVerticalSpacing(8)

        self.lang_box = QComboBox()
        self.lang_box.addItems(["ru", "en"])
        self.lang_box.currentTextChanged.connect(self.on_lang_change)
        self.runtime_box = QComboBox()
        self.runtime_box.addItems(["local", "cloud"])
        self.runtime_box.currentTextChanged.connect(self._runtime_changed)
        self.runtime_box.hide()
        self.local_btn = QPushButton("Local")
        self.local_btn.setCheckable(True)
        self.local_btn.setChecked(True)
        self.local_btn.setObjectName("SegmentBtn")
        self.cloud_btn = QPushButton("Cloud")
        self.cloud_btn.setCheckable(True)
        self.cloud_btn.setObjectName("SegmentBtn")
        self.local_btn.clicked.connect(lambda: self._set_runtime("local"))
        self.cloud_btn.clicked.connect(lambda: self._set_runtime("cloud"))
        self.model_box = QComboBox()
        self.model_box.addItems(self.catalog.names())
        self.model_box.currentTextChanged.connect(self._refresh_model_info)
        self.src_lang = QComboBox()
        self.src_lang.addItems(["auto", "en", "ru", "de", "fr", "es", "uk", "zh"])
        self.src_lang.setCurrentText("en")
        self.dst_lang = QComboBox()
        self.dst_lang.addItems(["ru", "en", "de", "fr", "es", "uk", "zh"])
        self.dst_lang.setCurrentText("ru")
        self.file_btn = QPushButton(self.tr("choose_file"))
        self.file_btn.clicked.connect(self.choose_file)
        self.translate_btn = QPushButton(self.tr("translate"))
        self.translate_btn.setObjectName("PrimaryButton")
        self.translate_btn.clicked.connect(self.translate_current)
        self.save_btn = QPushButton(self.tr("save_as"))
        self.save_btn.clicked.connect(self.save_result)

        grid.addWidget(QLabel("UI"), 0, 0)
        grid.addWidget(self.lang_box, 0, 1)
        grid.addWidget(QLabel(self.tr("runtime")), 0, 2)
        runtime_wrap = QFrame()
        runtime_wrap.setObjectName("SegmentWrap")
        runtime_layout = QHBoxLayout(runtime_wrap)
        runtime_layout.setContentsMargins(3, 3, 3, 3)
        runtime_layout.setSpacing(4)
        runtime_layout.addWidget(self.local_btn)
        runtime_layout.addWidget(self.cloud_btn)
        grid.addWidget(runtime_wrap, 0, 3)
        grid.addWidget(QLabel(self.tr("model")), 0, 4)
        grid.addWidget(self.model_box, 0, 5, 1, 3)
        self.swap_btn = QPushButton("⇄")
        self.swap_btn.setObjectName("SwapButton")
        self.swap_btn.clicked.connect(self._swap_langs)
        grid.addWidget(QLabel("Исходный язык"), 1, 0)
        grid.addWidget(self.src_lang, 1, 1)
        grid.addWidget(self.swap_btn, 1, 2)
        grid.addWidget(QLabel("Язык перевода"), 1, 3)
        grid.addWidget(self.dst_lang, 1, 4)
        grid.addWidget(self.file_btn, 1, 5)
        grid.addWidget(self.translate_btn, 1, 6)
        grid.addWidget(self.save_btn, 1, 7)
        page.addWidget(controls)

        self.cloud_group = QGroupBox("Cloud")
        cloud_form = QFormLayout(self.cloud_group)
        self.base_url_edit = QLineEdit("https://api.deepseek.com")
        self.api_key_edit = QLineEdit()
        self.api_key_edit.setEchoMode(QLineEdit.Password)
        cloud_form.addRow("Base URL", self.base_url_edit)
        cloud_form.addRow("API Key", self.api_key_edit)
        page.addWidget(self.cloud_group)

        info = QFrame()
        info.setObjectName("Card")
        info_layout = QVBoxLayout(info)
        info_layout.setContentsMargins(12, 8, 12, 8)
        self.compat_label = QLabel()
        self.compat_label.setWordWrap(True)
        info_layout.addWidget(self.compat_label)
        page.addWidget(info)

        split = QSplitter(Qt.Horizontal)
        left_panel = QFrame()
        left_panel.setObjectName("Card")
        left_layout = QVBoxLayout(left_panel)
        left_layout.setContentsMargins(8, 8, 8, 8)
        left_layout.addWidget(QLabel("Просмотр исходного файла"))
        left_inner = QHBoxLayout()
        self.source_pages = QListWidget()
        self.source_pages.setFixedWidth(92)
        self.source_box = QPlainTextEdit()
        self.source_box.setPlaceholderText(self.tr("source"))
        left_inner.addWidget(self.source_pages)
        left_inner.addWidget(self.source_box, 1)
        left_layout.addLayout(left_inner)

        right_panel = QFrame()
        right_panel.setObjectName("Card")
        right_layout = QVBoxLayout(right_panel)
        right_layout.setContentsMargins(8, 8, 8, 8)
        right_layout.addWidget(QLabel("Просмотр переведенного файла"))
        right_inner = QHBoxLayout()
        self.result_pages = QListWidget()
        self.result_pages.setFixedWidth(92)
        self.result_box = QPlainTextEdit()
        self.result_box.setPlaceholderText(self.tr("result"))
        right_inner.addWidget(self.result_pages)
        right_inner.addWidget(self.result_box, 1)
        right_layout.addLayout(right_inner)

        split.addWidget(left_panel)
        split.addWidget(right_panel)
        split.setSizes([700, 700])
        page.addWidget(split, 1)

        bottom = QHBoxLayout()
        self.pdf_to_docx = QCheckBox("PDF -> DOCX output")
        self.pdf_to_docx.setChecked(True)
        self.progress = QProgressBar()
        self.progress.setRange(0, 100)
        self.progress.setValue(0)
        bottom.addWidget(self.pdf_to_docx)
        bottom.addWidget(self.progress, 1)
        page.addLayout(bottom)

        self.status = QLabel()
        page.addWidget(self.status)
        self._runtime_changed("local")

    def _apply_styles(self) -> None:
        self.setStyleSheet(
            """
            QWidget { background: #f3f5f9; color: #1f2b3f; font-family: Segoe UI; font-size: 13px; }
            #HeaderCard { background: #ffffff; border: 1px solid #d9e1ef; border-radius: 14px; }
            #Card { background: #ffffff; border: 1px solid #d9e1ef; border-radius: 14px; }
            #Title { font-size: 34px; font-weight: 700; color: #1f2b3f; }
            #Subtitle { color: #6a7892; font-size: 14px; }
            #Badge { background: #eef8ef; color: #295f31; border: 1px solid #bfe0c4; border-radius: 10px; padding: 6px 10px; font-weight: 600; }
            #SegmentWrap { background: #f0f4fb; border: 1px solid #cad7ec; border-radius: 10px; }
            #SegmentBtn {
                background: transparent;
                border: none;
                border-radius: 8px;
                padding: 6px 12px;
                color: #486180;
                font-weight: 600;
            }
            #SegmentBtn:checked {
                background: #e6efff;
                color: #224d93;
                border: 1px solid #8fb0e5;
            }
            #SwapButton {
                min-width: 34px;
                max-width: 34px;
                min-height: 32px;
                max-height: 32px;
                background: #ffffff;
                border: 1px solid #cfd9eb;
                border-radius: 16px;
                padding: 0px;
                font-size: 15px;
                color: #5c6f8f;
            }
            #SwapButton:hover { background: #eef4ff; border-color: #98b5e8; color: #2b5ea9; }
            QGroupBox { background: #ffffff; border: 1px solid #d9e1ef; border-radius: 14px; margin-top: 10px; font-weight: 600; }
            QGroupBox::title { left: 12px; top: -3px; color: #6a7892; }
            QLineEdit, QComboBox, QPlainTextEdit {
                background: #fbfcfe; border: 1px solid #cfd9eb; border-radius: 10px; padding: 6px 10px; color: #1f2b3f;
            }
            QListWidget {
                background: #f7f9fd;
                border: 1px solid #cfd9eb;
                border-radius: 10px;
                padding: 4px;
            }
            QListWidget::item {
                border: 1px solid #d5e0f2;
                border-radius: 8px;
                margin: 3px;
                padding: 6px 5px;
                min-height: 28px;
            }
            QListWidget::item:selected {
                background: #e8f1ff;
                border-color: #7da7ef;
                color: #1f3e72;
            }
            QLineEdit:focus, QComboBox:focus, QPlainTextEdit:focus { border: 1px solid #4f86eb; }
            QPushButton {
                background: #ffffff; border: 1px solid #cfd9eb; border-radius: 10px; padding: 9px 14px; font-weight: 600; color: #2d3a52;
            }
            QPushButton:hover { background: #f4f7fd; }
            QPushButton:disabled { color: #8c97ac; background: #f2f5fa; border-color: #d6deec; }
            #PrimaryButton { background: #2f78ee; color: #ffffff; border-color: #2f78ee; }
            #PrimaryButton:hover { background: #2769d3; }
            QProgressBar { border: 1px solid #d0daec; border-radius: 8px; background: #edf2fa; text-align: center; color: #52627d; }
            QProgressBar::chunk { background: #2f78ee; border-radius: 7px; }
            """
        )

    def _runtime_changed(self, val: str) -> None:
        self.cloud_group.setVisible(val == "cloud")

    def _set_runtime(self, runtime: str) -> None:
        self.local_btn.setChecked(runtime == "local")
        self.cloud_btn.setChecked(runtime == "cloud")
        self.runtime_box.setCurrentText(runtime)
        self._runtime_changed(runtime)

    def _swap_langs(self) -> None:
        src = self.src_lang.currentText()
        dst = self.dst_lang.currentText()
        if src == "auto":
            return
        self.src_lang.setCurrentText(dst)
        self.dst_lang.setCurrentText(src)

    def on_lang_change(self, lang: str) -> None:
        self.lang = lang
        self.setWindowTitle(self.tr("title"))
        self.file_btn.setText(self.tr("choose_file"))
        self.translate_btn.setText(self.tr("translate"))
        self.save_btn.setText(self.tr("save_as"))

    def set_status(self, text: str) -> None:
        self.status.setText(text)

    def _set_busy(self, busy: bool) -> None:
        self.translate_btn.setEnabled(not busy)
        self.file_btn.setEnabled(not busy)
        self.save_btn.setEnabled(not busy)
        self.model_box.setEnabled(not busy)

    def _refresh_model_info(self) -> None:
        model = self.catalog.get(self.model_box.currentText())
        req_ram = model.get("required_ram_gb", 0)
        req_vram = model.get("required_vram_gb")
        ok_ram = self.specs.ram_gb >= req_ram
        ok_vram = True if req_vram is None else (self.specs.vram_gb or 0) >= req_vram
        icon = "OK" if (ok_ram and ok_vram) else "WARN"
        self.compat_label.setText(
            f"{self.tr('compat')}: {icon}\nCPU: {self.specs.cpu_name} ({self.specs.cpu_cores} cores)\n"
            f"RAM: {self.specs.ram_gb} GB (need {req_ram} GB)\n"
            f"VRAM: {self.specs.vram_gb if self.specs.vram_gb is not None else 'N/A'} GB (need {req_vram if req_vram is not None else 'N/A'} GB)\n"
            f"{model.get('description', '')}"
        )
        self.hw_badge.setText(
            f"RAM: {self.specs.ram_gb} / {req_ram} GB\n"
            f"VRAM: {self.specs.vram_gb if self.specs.vram_gb is not None else 'N/A'} / {req_vram if req_vram is not None else 'N/A'} GB"
        )

    def choose_file(self) -> None:
        path_str, _ = QFileDialog.getOpenFileName(self, self.tr("choose_file"), str(Path.home()), "Documents (*.txt *.md *.docx *.pdf);;All Files (*.*)")
        if not path_str:
            return
        self.current_file = Path(path_str)
        suffix = self.current_file.suffix.lower()
        if suffix in {".txt", ".md"}:
            text = self.current_file.read_text(encoding="utf-8", errors="ignore")
            self.page_count = 1
        elif suffix == ".docx":
            text = "\n".join(p.text for p in Document(self.current_file).paragraphs)
            self.page_count = max(1, min(12, len(Document(self.current_file).paragraphs) // 8 + 1))
        elif suffix == ".pdf":
            reader = PdfReader(str(self.current_file))
            self.page_count = len(reader.pages)
            text = "\n\n".join((p.extract_text() or "") for p in reader.pages).strip()
        else:
            QMessageBox.warning(self, APP_NAME, "Unsupported file type")
            return
        self.source_box.setPlainText(text)
        self._fill_page_lists()
        self.set_status(f"Loaded: {self.current_file.name}")

    def _fill_page_lists(self) -> None:
        self.source_pages.clear()
        self.result_pages.clear()
        count = max(1, min(self.page_count or 1, 12))
        for i in range(1, count + 1):
            src_item = QListWidgetItem(str(i))
            dst_item = QListWidgetItem(str(i))
            self.source_pages.addItem(src_item)
            self.result_pages.addItem(dst_item)
        if self.source_pages.count() > 0:
            self.source_pages.setCurrentRow(0)
            self.result_pages.setCurrentRow(0)

    def _make_client(self) -> AIClient:
        runtime = self.runtime_box.currentText()
        return AIClient(runtime, self.base_url_edit.text().strip() if runtime == "cloud" else "", self.api_key_edit.text().strip() if runtime == "cloud" else "")

    def _translate_task(self, model: str, text: str, src: str, dst: str, progress_callback: Signal) -> str:
        engine = TranslatorEngine(self._make_client())
        return engine.translate_text(model, text, src, dst, progress_cb=lambda i, t: progress_callback.emit((i, t)))

    def _on_translate_progress(self, payload: object) -> None:
        i, total = payload  # type: ignore[misc]
        self.progress.setValue(int((i / total) * 100))
        self.set_status(f"Translating chunk {i}/{total}...")

    def translate_current(self) -> None:
        source = self.source_box.toPlainText().strip()
        if not source:
            QMessageBox.warning(self, APP_NAME, "Load file or input text first")
            return
        worker = Worker(
            self._translate_task,
            self.model_box.currentText(),
            source,
            self.src_lang.text().strip() or "auto",
            self.dst_lang.text().strip() or "en",
        )
        self._set_busy(True)
        self.progress.setValue(0)
        self.set_status("Starting translation...")
        worker.signals.progress.connect(self._on_translate_progress)
        worker.signals.result.connect(self._on_translate_result)
        worker.signals.error.connect(lambda msg: QMessageBox.critical(self, APP_NAME, msg))
        worker.signals.finished.connect(lambda: self._set_busy(False))
        self.thread_pool.start(worker)

    def _on_translate_result(self, text: object) -> None:
        self.result_text = str(text)
        self.result_box.setPlainText(self.result_text)
        if self.result_pages.count() == 0:
            self._fill_page_lists()
        self.progress.setValue(100)
        self.set_status("Translation complete")

    def save_result(self) -> None:
        text = self.result_box.toPlainText().strip()
        if not text:
            QMessageBox.warning(self, APP_NAME, "No translation to save")
            return
        if self.current_file and self.current_file.suffix.lower() == ".pdf" and self.pdf_to_docx.isChecked():
            out, _ = QFileDialog.getSaveFileName(self, self.tr("save_as"), str(self.current_file.with_name(self.current_file.stem + "_translated.docx")), "Word (*.docx)")
            if out:
                doc = Document()
                for line in text.splitlines():
                    doc.add_paragraph(line)
                doc.save(out)
                self.set_status(f"Saved: {out}")
            return
        ext_filter = "Text (*.txt);;Markdown (*.md);;Word (*.docx)"
        out, _ = QFileDialog.getSaveFileName(self, self.tr("save_as"), str(Path.home() / "translated.txt"), ext_filter)
        if not out:
            return
        if out.lower().endswith(".docx"):
            doc = Document()
            for line in text.splitlines():
                doc.add_paragraph(line)
            doc.save(out)
        else:
            Path(out).write_text(text, encoding="utf-8")
        self.set_status(f"Saved: {out}")


def main() -> None:
    app = QApplication(sys.argv)
    app.setApplicationName(APP_NAME)
    w = MainWindow()
    w.show()
    sys.exit(app.exec())


if __name__ == "__main__":
    main()
